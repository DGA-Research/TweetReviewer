import os
import re
import base64
import tempfile
from datetime import datetime
from pathlib import Path
from io import BytesIO

import pandas as pd
import streamlit as st
import requests
import streamlit.components.v1 as components
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

WORD_FILENAME = "Issue Clipbook.docx"
SAVE_INTERVAL = 20

# Root for per-session ephemeral working files. In Docker this is a tmpfs mount
# (RAM-backed, wiped on restart) so nothing persists on the host/VPS.
SESSIONS_ROOT = os.environ.get("SESSIONS_ROOT", "/app/.sessions")


def get_session_dir() -> str:
    """Return a per-Streamlit-session temp directory for working files.

    Each browser session gets its own directory so concurrent researchers never
    overwrite each other's uploaded file, working copy, or Word clipbook.
    """
    session_dir = st.session_state.get("session_dir")
    if session_dir and os.path.isdir(session_dir):
        return session_dir
    try:
        os.makedirs(SESSIONS_ROOT, exist_ok=True)
        session_dir = tempfile.mkdtemp(prefix="sess_", dir=SESSIONS_ROOT)
    except Exception:
        # Fall back to the system temp dir (e.g. running locally without /app).
        session_dir = tempfile.mkdtemp(prefix="tweetrev_sess_")
    st.session_state.session_dir = session_dir
    return session_dir


def get_word_path() -> str:
    return os.path.join(get_session_dir(), WORD_FILENAME)


def save_working_copy(df: pd.DataFrame, path: str) -> None:
    """Persist the working dataframe back to its session-local file, matching format."""
    if str(path).lower().endswith(".csv"):
        df.to_csv(path, index=False)
    else:
        df.to_excel(path, index=False)


def auto_detect_column_mapping(df_columns: list[str]) -> dict[str, str]:
    """Auto-detect column mappings based on common patterns."""
    columns_lower = {col.lower(): col for col in df_columns}
    mapping = {}

    # URL mapping
    for pattern in ['url']:
        if pattern in columns_lower:
            mapping['url'] = columns_lower[pattern]
            break

    # Text mapping
    for pattern in ['text', 'content', 'tweet', 'message']:
        if pattern in columns_lower:
            mapping['text'] = columns_lower[pattern]
            break

    # Date mapping
    for pattern in ['date correct format', 'date', 'createdat', 'created_at', 'timestamp', 'posted_at']:
        if pattern in columns_lower:
            mapping['date'] = columns_lower[pattern]
            break

    # Quote tweet mapping
    for pattern in ['is_quote_tweet', 'isquote', 'is_quote', 'quote tweet', 'quote_tweet']:
        if pattern in columns_lower:
            mapping['quote'] = columns_lower[pattern]
            break

    # Bad words mapping (optional)
    for pattern in ['bad_words_found', 'flags', 'warnings']:
        if pattern in columns_lower:
            mapping['bad_words'] = columns_lower[pattern]
            break

    return mapping


def get_column_mapping(key: str, default: str = '') -> str:
    """Get the actual column name from the mapping."""
    if 'column_mapping' not in st.session_state:
        return default
    return st.session_state.column_mapping.get(key, default)


def trigger_rerun() -> None:
    rerun_fn = getattr(st, 'experimental_rerun', None) or getattr(st, 'rerun', None)
    if rerun_fn is None:
        raise RuntimeError('Streamlit rerun function is unavailable in this environment.')
    rerun_fn()





def ensure_hotkeys_script() -> None:
    if st.session_state.get("hotkeys_injected"):
        return

    components.html(
        """
        <script>
        (function() {
            const doc = window.parent.document;
            if (doc.hotkeysRegistered) { return; }
            doc.hotkeysRegistered = true;
            doc.addEventListener('keydown', function(event) {
                const active = doc.activeElement;
                if (active && ['INPUT', 'TEXTAREA'].includes(active.tagName)) {
                    return;
                }
                const key = event.key.toLowerCase();
                if (key === 'q' || key === 'w') {
                    const label = key === 'q' ? 'Pass' : 'Bullet';
                    const buttons = Array.from(doc.querySelectorAll('button'));
                    const target = buttons.find(btn => {
                        const text = (btn.innerText || '').trim();
                        return text === label;
                    });
                    if (target) {
                        target.click();
                        event.preventDefault();
                    }
                }
            }, true);
        })();
        </script>
        """,
        height=0,
        width=0,
    )
    st.session_state.hotkeys_injected = True












def extract_handle_from_url(url: str) -> str:
    if not isinstance(url, str):
        return 'unknown'

    match = re.search(r"https?://(?:www\.)?(?:twitter|x)\.com/([^/]+)", url, re.IGNORECASE)
    if not match:
        return 'unknown'

    handle = match.group(1).strip('@').strip()
    return handle or 'unknown'


def derive_export_metadata(df: pd.DataFrame) -> tuple[str, str, str, str]:
    handle = 'unknown'
    url_series = df.get('URL')
    if url_series is not None and not url_series.dropna().empty:
        handle = extract_handle_from_url(url_series.dropna().iloc[0])

    date_series = df.get('Date Correct Format')
    fallback_date_series = df.get('Date')
    first_date = 'unknown'
    last_tweet_date = 'unknown'
    if date_series is None or date_series.dropna().empty:
        date_series = fallback_date_series
    if date_series is not None and not date_series.dropna().empty:
        try:
            dates = pd.to_datetime(date_series.dropna(), utc=True)
            if not dates.empty:
                first_date = dates.min().strftime('%Y%m%d')
                last_tweet_date = dates.max().strftime('%Y%m%d')
        except Exception:
            pass

    last_reviewed_date = 'unknown'
    reviewed_series = df.get('Reviewed')
    if reviewed_series is not None:
        reviewed_mask = reviewed_series.fillna(False).astype(bool)
        if reviewed_mask.any():
            candidate_dates = df.loc[reviewed_mask, 'Date Correct Format'] if 'Date Correct Format' in df.columns else None
            if candidate_dates is None or candidate_dates.dropna().empty:
                candidate_dates = df.loc[reviewed_mask, 'Date'] if 'Date' in df.columns else None
            if candidate_dates is not None and not candidate_dates.dropna().empty:
                try:
                    reviewed_dates = pd.to_datetime(candidate_dates.dropna())
                    if not reviewed_dates.empty:
                        last_reviewed_date = reviewed_dates.max().strftime('%Y%m%d')
                except Exception:
                    pass

    return handle, first_date, last_reviewed_date, last_tweet_date


def build_export_filename(df: pd.DataFrame) -> str:
    handle, first_date, last_reviewed_date, last_tweet_date = derive_export_metadata(df)
    reviewed_series = df.get('Reviewed')
    any_reviewed = False
    if reviewed_series is not None:
        any_reviewed = reviewed_series.fillna(False).astype(bool).any()
    if any_reviewed:
        today = datetime.now().strftime('%Y%m%d')
        parts = ['REVIEWED', handle, first_date, last_reviewed_date, today]
    else:
        parts = ['UNREVIEWED', handle, first_date, last_tweet_date]
    safe_parts = [re.sub(r'[^A-Za-z0-9_-]+', '_', part) if part else 'unknown' for part in parts]
    return '_'.join(safe_parts) + '.xlsx'





def parse_review_filename(name: str) -> tuple[str, str, str, str, bool, str] | None:
    path = Path(name)
    suffix = path.suffix
    stem = path.stem
    is_auto = False
    if stem.endswith('_autoPush'):
        is_auto = True
        stem = stem[:-len('_autoPush')]
    parts = stem.split('_')
    if len(parts) < 4 or parts[0] != 'REVIEWED':
        return None
    today = parts[-1]
    last_date = parts[-2]
    first_date = parts[-3]
    handle = '_'.join(parts[1:-3]) or 'unknown'
    if not (len(first_date) == 8 and first_date.isdigit() and len(last_date) == 8 and last_date.isdigit() and len(today) == 8 and today.isdigit()):
        return None
    return handle, first_date, last_date, today, is_auto, suffix


def refresh_export_name() -> None:
    if 'df' not in st.session_state:
        return
    previous_auto = st.session_state.get('initial_export_name')
    new_name = build_export_filename(st.session_state.df)
    st.session_state.initial_export_name = new_name
    if 'export_name' not in st.session_state or st.session_state.get('export_name') == previous_auto:
        st.session_state.pending_export_name = new_name



def get_github_config() -> tuple[bool, dict | str]:
    """Resolve GitHub config from environment variables first (Docker/VPS),
    falling back to Streamlit secrets (Streamlit Cloud compatibility)."""
    secrets_cfg = {}
    if hasattr(st, 'secrets'):
        try:
            secrets_cfg = dict(st.secrets.get('github', {}))
        except Exception:
            secrets_cfg = {}

    def pick(env_key: str, secret_key: str, default=None):
        value = os.environ.get(env_key)
        if value is None or value == '':
            value = secrets_cfg.get(secret_key)
        return value if value not in (None, '') else default

    token = pick('GITHUB_TOKEN', 'token')
    owner = pick('GITHUB_OWNER', 'owner')
    repo = pick('GITHUB_REPO', 'repo')
    branch = pick('GITHUB_BRANCH', 'branch', 'main')
    # reviews_dir keeps the legacy 'target_dir' name for backward compatibility
    # with prune logic; inputs_dir is where researchers' raw uploads are stored.
    reviews_dir = (pick('GITHUB_REVIEWS_DIR', 'target_dir', '') or '').strip('/')
    inputs_dir = (pick('GITHUB_INPUTS_DIR', 'inputs_dir', '') or '').strip('/')

    if not token or not owner or not repo:
        return False, (
            'Missing GitHub config: set GITHUB_TOKEN / GITHUB_OWNER / GITHUB_REPO '
            '(environment) or github.token/owner/repo (st.secrets).'
        )

    return True, {
        'token': token,
        'owner': owner,
        'repo': repo,
        'branch': branch,
        'target_dir': reviews_dir,
        'reviews_dir': reviews_dir,
        'inputs_dir': inputs_dir,
    }


def github_headers(cfg: dict) -> dict:
    return {
        'Authorization': f"Bearer {cfg['token']}",
        'Accept': 'application/vnd.github+json',
        'X-GitHub-Api-Version': '2022-11-28',
    }


def github_put_bytes(cfg: dict, relative_path: str, content_bytes: bytes, message: str) -> tuple[bool, str]:
    """Create or update a file in the repo via the GitHub Contents API."""
    encoded_content = base64.b64encode(content_bytes).decode('utf-8')
    url = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents/{relative_path}"
    headers = github_headers(cfg)

    existing_sha = None
    get_response = requests.get(url, headers=headers, params={'ref': cfg['branch']})
    if get_response.status_code == 200:
        existing_sha = get_response.json().get('sha')
    elif get_response.status_code not in (200, 404):
        return False, f"GitHub API error ({get_response.status_code}): {get_response.text}"

    payload = {
        'message': message,
        'content': encoded_content,
        'branch': cfg['branch'],
    }
    if existing_sha:
        payload['sha'] = existing_sha

    put_response = requests.put(url, headers=headers, json=payload)
    if put_response.status_code not in (200, 201):
        return False, f"GitHub API error ({put_response.status_code}): {put_response.text}"

    return True, f"Uploaded and committed {Path(relative_path).name} to GitHub"


def github_list_dir(cfg: dict, directory: str) -> list[dict]:
    """List the entries of a repo directory; returns [] on any error."""
    directory = (directory or '').strip('/')
    base = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents"
    url = base if not directory else f"{base}/{directory}"
    try:
        resp = requests.get(url, headers=github_headers(cfg), params={'ref': cfg['branch']})
    except Exception:
        return []
    if resp.status_code != 200:
        return []
    try:
        entries = resp.json()
    except ValueError:
        return []
    return entries if isinstance(entries, list) else []


def github_get_file_bytes(cfg: dict, relative_path: str) -> bytes | None:
    """Download a single file's raw bytes from the repo."""
    url = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents/{relative_path}"
    try:
        resp = requests.get(url, headers=github_headers(cfg), params={'ref': cfg['branch']})
    except Exception:
        return None
    if resp.status_code != 200:
        return None
    try:
        data = resp.json()
    except ValueError:
        return None
    content = data.get('content')
    if content and data.get('encoding') == 'base64':
        try:
            return base64.b64decode(content)
        except Exception:
            pass
    # Large files (>1MB) come without inline content; use the download URL.
    download_url = data.get('download_url')
    if download_url:
        try:
            dl = requests.get(download_url, headers=github_headers(cfg))
            if dl.status_code == 200:
                return dl.content
        except Exception:
            return None
    return None


def push_input_to_github(local_path: str, original_name: str) -> tuple[bool, str]:
    """Push a researcher's raw input file (xlsx/csv) to the repo's inputs/ folder."""
    ok, cfg_or_message = get_github_config()
    if not ok:
        return False, cfg_or_message
    cfg = cfg_or_message
    try:
        content = Path(local_path).read_bytes()
    except Exception as exc:
        return False, f"Failed to read input file: {exc}"
    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", original_name) or "input"
    inputs_dir = cfg['inputs_dir']
    relative_path = safe_name if not inputs_dir else f"{inputs_dir}/{safe_name}"
    return github_put_bytes(cfg, relative_path, content, f"Add input file {safe_name}")


def save_and_git_commit(destination: Path, df: pd.DataFrame) -> tuple[bool, str]:
    try:
        destination.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(destination, index=False)
    except Exception as exc:
        return False, f"Failed to save workbook: {exc}"

    try:
        file_bytes = destination.read_bytes()
    except Exception as exc:
        return False, f"Failed to read saved workbook: {exc}"

    ok, cfg_or_message = get_github_config()
    if not ok:
        return False, cfg_or_message

    cfg = cfg_or_message
    relative_path = destination.name if not cfg['reviews_dir'] else f"{cfg['reviews_dir']}/{destination.name}"
    return github_put_bytes(cfg, relative_path, file_bytes, f"Add reviewed tweets {destination.name}")


def prune_previous_auto_push_files(current_destination: Path) -> None:
    parsed = parse_review_filename(current_destination.name)
    if not parsed:
        return
    handle, first_date, last_date, today, is_auto, suffix = parsed
    if not is_auto:
        return

    ok, cfg_or_message = get_github_config()
    if not ok:
        return
    cfg = cfg_or_message
    headers = {
        'Authorization': f"Bearer {cfg['token']}",
        'Accept': 'application/vnd.github+json',
        'X-GitHub-Api-Version': '2022-11-28',
    }
    base_contents_url = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents"
    target_dir = cfg['target_dir']
    list_url = base_contents_url if not target_dir else f"{base_contents_url}/{target_dir}"
    params = {'ref': cfg['branch']}

    try:
        list_response = requests.get(list_url, headers=headers, params=params)
    except Exception:
        return
    if list_response.status_code != 200:
        return
    try:
        entries = list_response.json()
    except ValueError:
        return
    if not isinstance(entries, list):
        return

    for entry in entries:
        if entry.get('type') != 'file':
            continue
        name = entry.get('name')
        parsed_entry = parse_review_filename(name)
        if not parsed_entry:
            continue
        e_handle, e_first, _e_last, _e_today, e_auto, _e_suffix = parsed_entry
        if not e_auto:
            continue
        if e_handle != handle or e_first != first_date:
            continue
        if name == current_destination.name:
            continue
        sha = entry.get('sha')
        if not sha:
            continue
        relative_name = name if not target_dir else f"{target_dir}/{name}"
        delete_url = f"{base_contents_url}/{relative_name}"
        payload = {
            'message': f"Remove previous auto push {name}",
            'sha': sha,
            'branch': cfg['branch'],
        }
        try:
            requests.delete(delete_url, headers=headers, json=payload)
        except Exception:
            continue



def prune_older_manual_reviews(current_destination: Path) -> None:
    parsed = parse_review_filename(current_destination.name)
    if not parsed:
        return
    handle, first_date, last_date, _today, is_auto, suffix = parsed
    if is_auto:
        return

    ok, cfg_or_message = get_github_config()
    if not ok:
        return
    cfg = cfg_or_message
    headers = {
        'Authorization': f"Bearer {cfg['token']}",
        'Accept': 'application/vnd.github+json',
        'X-GitHub-Api-Version': '2022-11-28',
    }
    base_contents_url = f"https://api.github.com/repos/{cfg['owner']}/{cfg['repo']}/contents"
    target_dir = cfg['target_dir']
    list_url = base_contents_url if not target_dir else f"{base_contents_url}/{target_dir}"
    params = {'ref': cfg['branch']}

    try:
        list_response = requests.get(list_url, headers=headers, params=params)
    except Exception:
        return
    if list_response.status_code != 200:
        return
    try:
        entries = list_response.json()
    except ValueError:
        return
    if not isinstance(entries, list):
        return

    for entry in entries:
        if entry.get('type') != 'file':
            continue
        name = entry.get('name')
        parsed_entry = parse_review_filename(name)
        if not parsed_entry:
            continue
        e_handle, e_first, e_last, _e_today, e_auto, _e_suffix = parsed_entry
        if e_auto:
            continue
        if e_handle != handle or e_first != first_date:
            continue
        if name == current_destination.name:
            continue
        if e_last >= last_date:
            continue
        sha = entry.get('sha')
        if not sha:
            continue
        relative_name = name if not target_dir else f"{target_dir}/{name}"
        delete_url = f"{base_contents_url}/{relative_name}"
        payload = {
            'message': f"Remove older review {name}",
            'sha': sha,
            'branch': cfg['branch'],
        }
        try:
            requests.delete(delete_url, headers=headers, json=payload)
        except Exception:
            continue



@st.cache_data(ttl=60, show_spinner=False)
def list_github_workbooks(owner: str, repo: str, branch: str, token: str,
                          inputs_dir: str, reviews_dir: str) -> list[dict]:
    """List .xlsx/.csv files available in the repo's inputs/ and reviews/ folders.

    Cached briefly to avoid hammering the GitHub API on every rerun; call
    ``list_github_workbooks.clear()`` after a push to refresh.
    """
    cfg = {'owner': owner, 'repo': repo, 'branch': branch, 'token': token}
    results: list[dict] = []
    seen: set[str] = set()
    for folder_label, directory in (('inputs', inputs_dir), ('reviews', reviews_dir)):
        for entry in github_list_dir(cfg, directory):
            if entry.get('type') != 'file':
                continue
            name = entry.get('name', '')
            if not name.lower().endswith(('.xlsx', '.csv')):
                continue
            path = entry.get('path') or (f"{directory}/{name}" if directory else name)
            if path in seen:
                continue
            seen.add(path)
            results.append({'label': f"{folder_label}/{name}", 'path': path, 'name': name})
    return sorted(results, key=lambda e: e['label'].lower())


def prepare_document(doc: Document) -> Document:
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)
    try:
        heading = doc.styles["Heading 2"]
    except KeyError:
        heading = doc.styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
    heading.font.name = "Arial"
    heading.font.size = Pt(12)
    heading.font.bold = True
    return doc


def load_dataframe(file_path: str, mapping_override: dict[str, str] | None = None) -> tuple[pd.DataFrame, int, int, dict[str, str], list[str]]:
    if str(file_path).lower().endswith(".csv"):
        raw_df = pd.read_csv(file_path)
    else:
        raw_df = pd.read_excel(file_path)
    raw_df.columns = raw_df.columns.str.strip()
    source_columns = list(raw_df.columns)

    # Auto-detect column mappings
    detected_mapping = auto_detect_column_mapping(source_columns)

    # Apply user overrides when provided
    overrides = mapping_override or {}
    for key, column_name in overrides.items():
        if column_name and column_name in source_columns:
            detected_mapping[key] = column_name

    # Guard against duplicate assignments
    selected_columns = [detected_mapping.get(k) for k in ['url', 'text', 'date', 'quote', 'bad_words']]
    selected_columns = [col for col in selected_columns if col]
    if len(selected_columns) != len(set(selected_columns)):
        raise ValueError("Duplicate column selections detected. Assign each input column to at most one field.")

    df = raw_df.copy()

    # Check for URL column (required)
    url_col = detected_mapping.get('url')
    if not url_col or url_col not in df.columns:
        raise ValueError(f"Expected a 'URL' column in the workbook. Available columns: {', '.join(source_columns)}")

    # Rename columns to standard names for internal use
    rename_map = {}
    if url_col and url_col != 'URL':
        rename_map[url_col] = 'URL'

    text_col = detected_mapping.get('text')
    if text_col and text_col != 'Text':
        rename_map[text_col] = 'Text'

    date_col = detected_mapping.get('date')
    if date_col and date_col not in ['Date', 'Date Correct Format']:
        rename_map[date_col] = 'Date'

    quote_col = detected_mapping.get('quote')
    if quote_col and quote_col != 'is_quote_tweet':
        rename_map[quote_col] = 'is_quote_tweet'

    bad_words_col = detected_mapping.get('bad_words')
    if bad_words_col and bad_words_col != 'bad_words_found':
        rename_map[bad_words_col] = 'bad_words_found'

    # Apply renames
    if rename_map:
        df = df.rename(columns=rename_map)

    original = len(df)
    df = df[df["URL"].fillna("").str.strip() != ""].reset_index(drop=True)
    removed_missing = original - len(df)

    before_dedup = len(df)
    df = df.drop_duplicates().reset_index(drop=True)
    removed_duplicates = before_dedup - len(df)

    if 'Date Correct Format' in df.columns:
        sort_series = pd.to_datetime(df['Date Correct Format'], errors='coerce', utc=True)
    else:
        sort_series = None
    if sort_series is None or sort_series.dropna().empty:
        sort_series = pd.to_datetime(df['Date'], errors='coerce', utc=True) if 'Date' in df.columns else None
    if sort_series is not None:
        df = df.assign(_sort_date=sort_series).sort_values('_sort_date', kind='stable', na_position='last').drop(columns='_sort_date').reset_index(drop=True)

    if 'Reviewed' not in df.columns:
        df['Reviewed'] = False
    df['Reviewed'] = df['Reviewed'].fillna(False).astype(bool)

    if 'Bullet topic' not in df.columns:
        df['Bullet topic'] = ''
    df['Bullet topic'] = df['Bullet topic'].fillna('').astype(str)

    return df, removed_missing, removed_duplicates, detected_mapping, source_columns

def rebuild_content_from_df(df: pd.DataFrame) -> None:
    st.session_state.content_by_topic = {}
    st.session_state.topic_history = []
    if 'Bullet topic' not in df.columns or 'Reviewed' not in df.columns:
        return
    reviewed_mask = df['Reviewed'].fillna(False).astype(bool)
    bulleted = df[reviewed_mask & df['Bullet topic'].fillna('').astype(str).str.strip().ne('')]
    for _, row in bulleted.iterrows():
        topic = str(row['Bullet topic']).strip().upper()
        format_text_for_bullet(row, topic)
        if topic not in st.session_state.topic_history:
            st.session_state.topic_history.append(topic)
            
def initialize_state(file_path: str, source_label: str | None = None, mapping_override: dict[str, str] | None = None) -> None:
    overrides_store = st.session_state.setdefault('column_mapping_overrides', {})
    active_override = {k: v for k, v in (mapping_override or {}).items() if v}
    df, removed_missing, removed_duplicates, detected_mapping, source_columns = load_dataframe(file_path, active_override or None)
    is_new_file = st.session_state.get('excel_path') != file_path
    st.session_state.df = df
    st.session_state.excel_path = file_path
    st.session_state.source_label = source_label or Path(file_path).name
    st.session_state.removed_rows_without_url = removed_missing
    st.session_state.removed_duplicate_rows = removed_duplicates
    st.session_state.removed_rows = removed_missing + removed_duplicates
    st.session_state.column_mapping = detected_mapping
    st.session_state.available_columns = source_columns
    overrides_store[file_path] = active_override
    st.session_state.original_columns = list(df.columns)
    if removed_missing or removed_duplicates:
        save_working_copy(df, file_path)

    word_path = get_word_path()
    if os.path.exists(word_path):
        doc = Document(word_path)
    else:
        doc = Document()
    st.session_state.doc = prepare_document(doc)

    st.session_state.topic_input = ''
    st.session_state.topic_select = ''
    st.session_state.clear_topic_inputs = False

    if is_new_file:
        st.session_state.content_by_topic = {}
        st.session_state.topic_history = []
        st.session_state.history_stack = []
        st.session_state.current_index = 0
        st.session_state.actions_since_save = 0
        st.session_state.last_save_message = None
        st.session_state.last_export_message = None
        st.session_state.last_export_success = None
        st.session_state.initial_export_name = build_export_filename(df)
        st.session_state.reset_export_name = False
        st.session_state.export_name = st.session_state.initial_export_name
        update_counts()
        refresh_export_name()
        advance_to_next_unreviewed()
        rebuild_content_from_df(df)

def update_counts() -> None:
    df = st.session_state.df
    reviewed = df['Reviewed'].fillna(False).astype(bool) if 'Reviewed' in df.columns else pd.Series(dtype=bool)
    bullet_topics = df['Bullet topic'].fillna('').astype(str).str.strip() if 'Bullet topic' in df.columns else pd.Series([''] * len(df))
    bullet_mask = reviewed & (bullet_topics != '')
    pass_mask = reviewed & ~bullet_mask
    st.session_state.pass_count = int(pass_mask.sum())
    st.session_state.bullet_count = int(bullet_mask.sum())
    st.session_state.total_reviewed = st.session_state.pass_count + st.session_state.bullet_count


def advance_to_next_unreviewed() -> None:
    df = st.session_state.df
    idx = st.session_state.current_index
    while idx < len(df) and bool(df.at[idx, 'Reviewed']):
        idx += 1
    st.session_state.current_index = idx


def save_progress(force: bool = False) -> None:
    if not force and st.session_state.actions_since_save < SAVE_INTERVAL:
        return
    save_working_copy(st.session_state.df, st.session_state.excel_path)
    st.session_state.doc.save(get_word_path())

    export_name = st.session_state.get('export_name')
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if export_name:
        base_destination = Path(export_name)
        auto_filename = f"{base_destination.stem}_autoPush{base_destination.suffix}" if base_destination.suffix else f"{base_destination.name}_autoPush"
        auto_destination = base_destination.with_name(auto_filename)
        if not auto_destination.is_absolute():
            auto_destination = Path(get_session_dir()) / auto_destination
        auto_destination.parent.mkdir(parents=True, exist_ok=True)
        success, message = save_and_git_commit(auto_destination, st.session_state.df)
        st.session_state.last_export_message = message
        st.session_state.last_export_success = success
        if success:
            prune_previous_auto_push_files(auto_destination)
        save_note = 'Progress auto-pushed' if success else 'Auto push failed'
    else:
        st.session_state.last_export_success = None
        st.session_state.last_export_message = None
        save_note = 'Progress auto-saved'
    st.session_state.actions_since_save = 0
    st.session_state.last_save_message = f"{save_note} at {timestamp}"


def increment_action_counter() -> None:
    st.session_state.actions_since_save += 1
    save_progress()


def normalize_spaces(text: str) -> str:
    text = re.sub(r"([.?!]) {2,}", r"\\1 ", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def add_hyperlink_date_only(paragraph, prefix: str, date_part: str, suffix: str, url: str) -> None:
    run_prefix = paragraph.add_run(prefix)
    run_prefix.font.name = "Arial"
    run_prefix.font.size = Pt(10)

    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    new_run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_props.append(run_style)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    new_run.append(run_props)
    text_element = OxmlElement("w:t")
    text_element.text = date_part
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    run_suffix = paragraph.add_run(suffix)
    run_suffix.font.name = "Arial"
    run_suffix.font.size = Pt(10)


def rebuild_document() -> None:
    doc = st.session_state.doc
    content = st.session_state.content_by_topic

    for paragraph in list(doc.paragraphs):
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    for topic in sorted(content.keys()):
        doc.add_paragraph(topic, style="Heading 2")
        for entry in content[topic]:
            para = doc.add_paragraph()
            run = para.add_run(entry["quoted_text"] + " ")
            run.font.name = "Arial"
            run.font.size = Pt(10)
            add_hyperlink_date_only(para, "[X, @RandyFeenstra, ", entry["date_str"], "]", entry["url"])

            doc.add_paragraph()
            centered = doc.add_paragraph()
            centered.alignment = 1
            add_hyperlink_date_only(centered, "[X, @RandyFeenstra, ", entry["date_str"], "]", entry["url"])
            doc.add_paragraph()


def format_text_for_bullet(row: pd.Series, topic: str) -> None:
    text = str(row.get("Text", ""))
    url = str(row.get("URL", ""))
    date_value = row.get("Date Correct Format")
    if pd.isna(date_value):
        date_value = row.get("Date")
    date = pd.to_datetime(date_value)
    date_str = f"{date.month}/{date.day}/{str(date.year)[2:]}"

    text = text.replace('"', "'").replace("\n", "\u00A0")
    text = normalize_spaces(text)
    quoted = f'"{text}"'

    entries = st.session_state.content_by_topic.setdefault(topic, [])
    entries.append({"quoted_text": quoted, "url": url, "date_str": date_str})
    rebuild_document()


def handle_pass() -> None:
    df = st.session_state.df
    idx = st.session_state.current_index
    prev_reviewed = bool(df.at[idx, 'Reviewed']) if 'Reviewed' in df.columns else False
    prev_topic_value = df.at[idx, 'Bullet topic'] if 'Bullet topic' in df.columns else ''
    prev_topic = '' if pd.isna(prev_topic_value) else str(prev_topic_value)
    df.at[idx, 'Reviewed'] = True
    df.at[idx, 'Bullet topic'] = ''
    st.session_state.history_stack.append({
        'index': idx,
        'action': 'pass',
        'prev_reviewed': prev_reviewed,
        'prev_topic': prev_topic,
    })
    st.session_state.current_index += 1
    update_counts()
    refresh_export_name()
    increment_action_counter()
    advance_to_next_unreviewed()


def handle_bullet(topic: str) -> None:
    df = st.session_state.df
    idx = st.session_state.current_index
    topic_clean = topic.strip()
    topic_upper = topic_clean.upper()
    prev_reviewed = bool(df.at[idx, 'Reviewed']) if 'Reviewed' in df.columns else False
    prev_topic_value = df.at[idx, 'Bullet topic'] if 'Bullet topic' in df.columns else ''
    prev_topic = '' if pd.isna(prev_topic_value) else str(prev_topic_value)
    df.at[idx, 'Reviewed'] = True
    df.at[idx, 'Bullet topic'] = topic_clean
    st.session_state.history_stack.append({
        'index': idx,
        'action': 'bullet',
        'prev_reviewed': prev_reviewed,
        'prev_topic': prev_topic,
        'topic': topic_upper,
    })
    format_text_for_bullet(df.iloc[idx], topic_upper)
    if topic_upper not in st.session_state.topic_history:
        st.session_state.topic_history.append(topic_upper)
    st.session_state.current_index += 1
    update_counts()
    refresh_export_name()
    increment_action_counter()
    advance_to_next_unreviewed()


def handle_back() -> bool:
    if not st.session_state.history_stack:
        return False

    entry = st.session_state.history_stack.pop()
    idx = entry.get('index')
    action = entry.get('action')
    topic = entry.get('topic')

    if action == 'bullet' and topic:
        entries = st.session_state.content_by_topic.get(topic, [])
        if entries:
            entries.pop()
            if not entries:
                st.session_state.content_by_topic.pop(topic, None)
        rebuild_document()

    prev_reviewed = entry.get('prev_reviewed', False)
    prev_topic = entry.get('prev_topic', '')
    st.session_state.df.at[idx, 'Reviewed'] = prev_reviewed
    st.session_state.df.at[idx, 'Bullet topic'] = prev_topic

    st.session_state.current_index = idx
    st.session_state.actions_since_save = max(st.session_state.actions_since_save - 1, 0)
    update_counts()
    refresh_export_name()
    advance_to_next_unreviewed()
    return True


def reset_for_rereview() -> None:
    st.session_state.df['Reviewed'] = False
    st.session_state.df['Bullet topic'] = ''
    st.session_state.history_stack = []
    st.session_state.content_by_topic = {}
    st.session_state.topic_history = []
    st.session_state.current_index = 0
    st.session_state.actions_since_save = 0
    st.session_state.clear_topic_inputs = False
    st.session_state.doc = prepare_document(Document())
    st.session_state.last_export_message = None
    st.session_state.last_export_success = None
    st.session_state.initial_export_name = build_export_filename(st.session_state.df)
    st.session_state.reset_export_name = True
    st.session_state.pop('export_name', None)
    update_counts()
    refresh_export_name()
    advance_to_next_unreviewed()
    save_working_copy(st.session_state.df, st.session_state.excel_path)
    save_progress(force=True)


def _check_password() -> bool:
    """Return True if the user has entered the correct app password.

    Reads STREAMLIT_PASSWORD from env (set in Coolify's env var panel).
    If the env var is not set, the gate is disabled and everyone can access the app.
    """
    required = os.environ.get("STREAMLIT_PASSWORD", "")
    if not required:
        return True

    if st.session_state.get("_authenticated"):
        return True

    st.set_page_config(page_title="Tweet Reviewer — Login", layout="centered")
    st.title("Tweet Reviewer")
    st.text_input("Password", type="password", key="_pw_input")
    if st.button("Log in"):
        if st.session_state.get("_pw_input") == required:
            st.session_state["_authenticated"] = True
            st.rerun()
        else:
            st.error("Incorrect password.")
    return False


def main() -> None:
    if not _check_password():
        st.stop()

    st.set_page_config(page_title="Tweet Reviewer", layout="wide")
    st.title("Tweet Reviewer")

    mapping_overrides = st.session_state.setdefault('column_mapping_overrides', {})

    uploaded_file = st.sidebar.file_uploader("Upload workbook (.xlsx or .csv)", type=["xlsx", "csv"])
    if uploaded_file is not None:
        uploaded_bytes = uploaded_file.getvalue()
        file_signature = (uploaded_file.name, len(uploaded_bytes))
        if st.session_state.get("uploaded_file_signature") != file_signature:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            src_suffix = Path(uploaded_file.name).suffix.lower()
            if src_suffix not in (".xlsx", ".csv"):
                src_suffix = ".xlsx"
            safe_stem = re.sub(r"[^A-Za-z0-9_-]+", "_", Path(uploaded_file.name).stem) or "workbook"
            destination = Path(get_session_dir()) / f"uploaded_{timestamp}_{safe_stem}{src_suffix}"
            destination.write_bytes(uploaded_bytes)
            st.session_state.uploaded_file_signature = file_signature
            # Remember the raw upload so the researcher can push it to inputs/.
            st.session_state.input_source_path = str(destination)
            st.session_state.input_source_name = uploaded_file.name
            initialize_state(str(destination), uploaded_file.name)
            trigger_rerun()
            st.stop()

    # GitHub-backed picker: load any .xlsx/.csv already stored in inputs/ or reviews/.
    ok_cfg, cfg_or_msg = get_github_config()
    with st.sidebar.expander("Load from GitHub", expanded=("df" not in st.session_state)):
        if not ok_cfg:
            st.caption("GitHub not configured; upload a file to get started.")
            st.caption(str(cfg_or_msg))
        else:
            cfg = cfg_or_msg
            choices = list_github_workbooks(
                cfg['owner'], cfg['repo'], cfg['branch'], cfg['token'],
                cfg['inputs_dir'], cfg['reviews_dir'],
            )
            if st.button("Refresh list", key="refresh_github_list"):
                list_github_workbooks.clear()
                trigger_rerun()
            if not choices:
                st.caption("No .xlsx/.csv files found in inputs/ or reviews/ yet.")
            else:
                labels = [c['label'] for c in choices]
                chosen_label = st.selectbox("Workbook on GitHub", labels, key="github_pick_label")
                if st.button("Load selected file", key="load_github_file"):
                    chosen = next((c for c in choices if c['label'] == chosen_label), None)
                    if chosen is not None:
                        data = github_get_file_bytes(cfg, chosen['path'])
                        if data is None:
                            st.error("Failed to download file from GitHub.")
                        else:
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", chosen['name'])
                            local_path = Path(get_session_dir()) / f"gh_{timestamp}_{safe_name}"
                            local_path.write_bytes(data)
                            # File already lives on GitHub; no raw-input push needed.
                            st.session_state.input_source_path = None
                            st.session_state.input_source_name = None
                            initialize_state(str(local_path), chosen['label'],
                                             mapping_overrides.get(str(local_path)))
                            trigger_rerun()
                            st.stop()

    if "df" not in st.session_state:
        st.warning("Upload a workbook or load one from GitHub to get started.")
        return

    # Display column mappings
    if st.session_state.get('column_mapping'):
        with st.sidebar.expander("Column Mappings", expanded=False):
            st.caption("Choose which workbook columns map to Tweet Reviewer fields.")
            available_columns = list(st.session_state.get('available_columns', []))
            mapping = st.session_state.column_mapping

            if not available_columns:
                st.info("No columns available in this workbook.")
            else:
                def option_index(options, value):
                    return options.index(value) if value in options else 0

                url_options = available_columns
                url_choice = st.selectbox(
                    "URL column",
                    url_options,
                    index=option_index(url_options, mapping.get('url')),
                )

                optional_options = [''] + available_columns
                text_choice = st.selectbox(
                    "Text column (optional)",
                    optional_options,
                    index=option_index(optional_options, mapping.get('text')),
                )
                date_choice = st.selectbox(
                    "Date column (optional)",
                    optional_options,
                    index=option_index(optional_options, mapping.get('date')),
                )
                quote_choice = st.selectbox(
                    "Quote flag column (optional)",
                    optional_options,
                    index=option_index(optional_options, mapping.get('quote')),
                )
                bad_words_choice = st.selectbox(
                    "Flags column (optional)",
                    optional_options,
                    index=option_index(optional_options, mapping.get('bad_words')),
                )

                action_cols = st.columns(2)
                with action_cols[0]:
                    apply_mapping = st.button("Apply mapping")
                with action_cols[1]:
                    reset_mapping = st.button("Use auto-detected mapping")

                if apply_mapping:
                    new_mapping = {
                        'url': url_choice,
                        'text': text_choice or '',
                        'date': date_choice or '',
                        'quote': quote_choice or '',
                        'bad_words': bad_words_choice or '',
                    }
                    selected_values = [col for col in new_mapping.values() if col]
                    duplicates = {col for col in selected_values if selected_values.count(col) > 1}
                    if duplicates:
                        st.warning(f"Columns used more than once: {', '.join(sorted(duplicates))}. Select each column only once.")
                    else:
                        cleaned_override = {k: v for k, v in new_mapping.items() if v}
                        st.session_state.column_mapping_overrides[st.session_state.excel_path] = cleaned_override
                        initialize_state(st.session_state.excel_path, st.session_state.source_label, cleaned_override)
                        trigger_rerun()
                        st.stop()

                if reset_mapping:
                    st.session_state.column_mapping_overrides[st.session_state.excel_path] = {}
                    initialize_state(st.session_state.excel_path, st.session_state.source_label)
                    trigger_rerun()
                    st.stop()

    st.sidebar.metric("Passed", st.session_state.pass_count)
    st.sidebar.metric("Bulleted", st.session_state.bullet_count)
    st.sidebar.metric("Total Reviewed", st.session_state.total_reviewed)

    if st.session_state.get("source_label"):
        st.sidebar.caption(f"Reviewing: {st.session_state.source_label}")

    removed_missing = st.session_state.get('removed_rows_without_url', 0)
    removed_duplicates = st.session_state.get('removed_duplicate_rows', 0)
    if removed_missing or removed_duplicates:
        parts = []
        if removed_missing:
            parts.append(f"{removed_missing} without a URL")
        if removed_duplicates:
            parts.append(f"{removed_duplicates} duplicate rows")
        st.sidebar.info("Removed " + " and ".join(parts))

    if st.session_state.last_save_message:
        st.sidebar.caption(st.session_state.last_save_message)

    if 'df' in st.session_state:
        refresh_export_name()
        if st.session_state.get('reset_export_name'):
            st.session_state.export_name = st.session_state.initial_export_name
            st.session_state.reset_export_name = False

        pending_name = st.session_state.pop('pending_export_name', None)
        if pending_name is not None:
            st.session_state.export_name = pending_name
        elif 'export_name' not in st.session_state:
            st.session_state.export_name = st.session_state.initial_export_name

        if 'export_name' in st.session_state:
            st.sidebar.text_input("Output xlsx filename:", value=st.session_state.export_name, key="export_name")

            if st.sidebar.button("Push xlsx to Git"):
                destination = Path(st.session_state.export_name)
                if not destination.is_absolute():
                    destination = Path(get_session_dir()) / destination
                destination.parent.mkdir(parents=True, exist_ok=True)
                success, message = save_and_git_commit(destination, st.session_state.df)
                st.session_state.last_export_message = message
                st.session_state.last_export_success = success
                if success:
                    prune_older_manual_reviews(destination)
                    st.session_state.initial_export_name = build_export_filename(st.session_state.df)
                    st.session_state.reset_export_name = True

            local_filename = st.session_state.export_name or "reviewed.xlsx"
            download_buffer = BytesIO()
            st.session_state.df.to_excel(download_buffer, index=False)
            download_buffer.seek(0)
            st.sidebar.download_button(
                "Download .xlsx sheet",
                data=download_buffer,
                file_name=local_filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_local_copy",
            )

            word_buffer = BytesIO()
            st.session_state.doc.save(word_buffer)
            word_buffer.seek(0)
            st.sidebar.download_button(
                "Download .docx summary",
                data=word_buffer,
                file_name=WORD_FILENAME,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word_copy",
            )

        if st.session_state.get("input_source_path"):
            if st.sidebar.button("Save input to GitHub (inputs/)", key="push_input_github"):
                ok_in, msg_in = push_input_to_github(
                    st.session_state["input_source_path"],
                    st.session_state.get("input_source_name", "input.xlsx"),
                )
                st.session_state.last_export_message = msg_in
                st.session_state.last_export_success = ok_in
                if ok_in:
                    list_github_workbooks.clear()
                trigger_rerun()
    if st.session_state.get("last_export_message"):
        message = st.session_state.last_export_message
        status = st.session_state.get("last_export_success")
        if status is True:
            st.sidebar.success(message)
        elif status is False:
            st.sidebar.error(message)
        else:
            st.sidebar.caption(message)

    if st.session_state.total_reviewed:
        st.sidebar.warning("Existing review marks detected.")
        if st.sidebar.button("Reset for re-review"):
            reset_for_rereview()
            trigger_rerun()

    advance_to_next_unreviewed()

    df = st.session_state.df
    idx = st.session_state.current_index

    if idx >= len(df):
        st.success("All rows reviewed.")
        return

    row = df.iloc[idx]
    st.subheader(f"Row {idx + 1} of {len(df)}")
    st.write(row.get("Text", ""))

    flag_value = row.get("bad_words_found", "")
    if pd.notna(flag_value) and str(flag_value).strip():
        st.warning(f"Flags: {flag_value}")

    quote_flag = row.get("is_quote_tweet") or row.get("Quote Tweet")
    if pd.notna(quote_flag) and quote_flag:
        st.info("QUOTE TWEET")

    url = row.get("URL", "")
    if isinstance(url, str) and url.strip():
        st.markdown(f"[Open Link]({url})")

    st.caption(f"Progress pushed to GitHub every {SAVE_INTERVAL} actions")

    if 'topic_input' not in st.session_state:
        st.session_state.topic_input = ''
    if 'topic_select' not in st.session_state:
        st.session_state.topic_select = ''
    if 'clear_topic_inputs' not in st.session_state:
        st.session_state.clear_topic_inputs = False
    if st.session_state.clear_topic_inputs:
        st.session_state.topic_input = ''
        st.session_state.topic_select = ''
        st.session_state.clear_topic_inputs = False

    columns = st.columns(3)
    ensure_hotkeys_script()
    if columns[0].button("Pass", key="pass_button"):
        handle_pass()
        trigger_rerun()

    with st.expander("Topic options", expanded=False):
        existing_topics = [''] + sorted(st.session_state.topic_history)
        st.selectbox("Choose existing topic", existing_topics, key="topic_select")
        st.text_input("Or enter a topic", key="topic_input")

    if columns[1].button("Bullet", key="bullet_button"):
        topic_choice = (st.session_state.topic_input or '').strip() or (st.session_state.topic_select or '').strip()
        if not topic_choice:
            st.warning("Provide a topic before marking as bullet.")
        else:
            handle_bullet(topic_choice)
            st.session_state.clear_topic_inputs = True
            trigger_rerun()

    if columns[2].button("Undo last"):
        if handle_back():
            trigger_rerun()
        else:
            st.info("Nothing to undo yet.")



if __name__ == "__main__":
    main()
